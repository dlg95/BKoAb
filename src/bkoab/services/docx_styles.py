from decimal import Decimal

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def format_eur(value: Decimal | float) -> str:
    amount = Decimal(str(value)).quantize(Decimal("0.01"))
    formatted = f"{amount:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"{formatted} €"


def _zero_paragraph_spacing(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def add_styled_paragraph(
    doc,
    text: str,
    *,
    bold: bool = False,
    size: int = 10,
    align: int | None = None,
    compact: bool = False,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if compact:
        _zero_paragraph_spacing(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def add_money_paragraph(
    doc,
    label: str,
    amount: Decimal | float,
    *,
    bold: bool = False,
    size: int = 10,
    compact: bool = False,
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if compact:
        _zero_paragraph_spacing(p)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = bold
    run_label.font.size = Pt(size)
    run_label.font.name = "Calibri"
    run_amount = p.add_run(format_eur(amount))
    run_amount.bold = bold
    run_amount.font.size = Pt(size)
    run_amount.font.name = "Calibri"
    return p


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 9,
    align_right: bool = False,
    shade: bool = False,
    compact: bool = False,
):
    cell.text = ""
    p = cell.paragraphs[0]
    if align_right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if compact:
        _zero_paragraph_spacing(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if shade:
        _set_cell_shading(cell, "F2F2F2")


def add_cell_paragraph(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 10,
    align_right: bool = False,
):
    p = cell.add_paragraph()
    if align_right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _zero_paragraph_spacing(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "CCCCCC")
        borders.append(element)
    tbl_pr.append(borders)


def set_table_no_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def set_cell_margins(cell, *, top: int = 20, bottom: int = 20, left: int = 40, right: int = 40) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_table_column_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)
