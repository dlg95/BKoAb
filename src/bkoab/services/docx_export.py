from dataclasses import dataclass
from datetime import date
from decimal import Decimal
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from bkoab.schemas import AllocationKey, PartySettlement, SettlementPreview
from bkoab.services.docx_styles import (
    add_cell_paragraph,
    add_styled_paragraph,
    format_eur,
    set_cell_margins,
    set_cell_text,
    set_table_column_widths,
    set_table_no_borders,
)


@dataclass
class PersonPeriodLine:
    valid_from: str
    valid_to: str | None
    persons: int


def _default_payment_text(apartment_iban: str, account_holder: str, reference: str) -> str:
    return (
        f"Bitte überweisen Sie den offenen Betrag auf folgendes Konto: "
        f"IBAN {apartment_iban}, Kontoinhaber {account_holder}"
        + (f", Verwendungszweck {reference}" if reference else "")
        + ". Ein Guthaben überweisen wir zeitnah auf Ihr uns bekanntes Konto."
    )


def _year_bounds(year: int) -> tuple[date, date]:
    return date(year, 1, 1), date(year, 12, 31)


def person_period_lines_for_year(
    person_period_lines: list[PersonPeriodLine],
    year: int,
) -> list[PersonPeriodLine]:
    """Clip person periods to the overlap with the billing calendar year."""
    year_start, year_end = _year_bounds(year)
    clipped: list[PersonPeriodLine] = []
    for period in person_period_lines:
        period_start = date.fromisoformat(period.valid_from)
        period_end = date.fromisoformat(period.valid_to) if period.valid_to else year_end
        start = max(period_start, year_start)
        end = min(period_end, year_end)
        if start > end:
            continue
        clipped.append(
            PersonPeriodLine(
                valid_from=start.isoformat(),
                valid_to=end.isoformat(),
                persons=period.persons,
            )
        )
    return clipped


def _format_person_periods_note(
    person_period_lines: list[PersonPeriodLine],
    year: int,
) -> str | None:
    year_periods = person_period_lines_for_year(person_period_lines, year)
    if not year_periods:
        return None
    periods_text = ", ".join(
        f"{period.persons} Pers. ({period.valid_from}–{period.valid_to})"
        for period in year_periods
    )
    return f"Personenzahl {year}: {periods_text}."


def _add_allocation_notes(
    doc: Document,
    *,
    preview: SettlementPreview,
    party: PartySettlement,
    person_period_lines: list[PersonPeriodLine],
) -> None:
    note = (
        f"Verteilung nach Personenmonaten: Gesamtkosten × (Ihre PM ÷ PM gesamt). "
        f"Ihre PM {party.head_months:.2f}, PM gesamt {preview.total_head_months:.2f}"
    )
    if float(preview.landlord_vacancy_head_months) > 0:
        note += f", Leerstand {preview.landlord_vacancy_head_months:.2f}"
    add_styled_paragraph(doc, note + ".", size=8, compact=True)

    if preview.unit_area_sqm and preview.total_property_area_sqm:
        add_styled_paragraph(
            doc,
            f"Flächenverteilung Gebäude: {preview.unit_area_sqm:.2f} m² von "
            f"{preview.total_property_area_sqm:.2f} m² gesamt.",
            size=8,
            compact=True,
        )

    periods_note = _format_person_periods_note(person_period_lines, preview.year)
    if periods_note:
        add_styled_paragraph(doc, periods_note, size=8, compact=True)


def _format_numerator(line, key: AllocationKey) -> str:
    if key == AllocationKey.PERSONENMONATE:
        return f"{line.party_numerator:.2f}"
    if line.party_denominator == line.party_numerator and line.party_denominator > 0:
        return "gleich"
    return f"{line.party_numerator:.2f}"


def _format_denominator(line, key: AllocationKey) -> str:
    if key == AllocationKey.PERSONENMONATE:
        return f"{line.party_denominator:.2f}"
    if line.party_denominator > 0 and line.party_numerator == Decimal("1"):
        return f"{line.party_denominator:.0f}"
    return f"{line.party_denominator:.2f}"


def _add_cost_lines_table(doc: Document, *, party: PartySettlement) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_no_borders(table)
    set_table_column_widths(table, [5.5, 2.2, 1.3, 1.3, 1.3, 2.0])

    headers = ["Kostenart", "Gesamt", "Quote", "Ihr Wert", "Gesamt", "Ihr Anteil"]
    for idx, title in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_margins(cell, top=10, bottom=20, left=20, right=20)
        set_cell_text(
            cell,
            title,
            bold=True,
            size=8,
            align_right=idx > 0,
            compact=True,
        )

    for line in party.cost_lines:
        key = line.allocation_key
        quote_label = "PM" if key == AllocationKey.PERSONENMONATE else "m²"
        row = table.add_row().cells
        values = [
            (line.label, False),
            (format_eur(line.total_prorated), True),
            (quote_label, True),
            (_format_numerator(line, key), True),
            (_format_denominator(line, key), True),
            (format_eur(line.party_share), True),
        ]
        for idx, (text, align_right) in enumerate(values):
            set_cell_margins(row[idx], top=8, bottom=8, left=20, right=20)
            set_cell_text(row[idx], text, size=8, align_right=align_right, compact=True)

    if any(line.has_document for line in party.cost_lines):
        add_styled_paragraph(
            doc,
            "Belege zu den markierten Kostenarten liegen der Abrechnung bei.",
            size=8,
            compact=True,
        )


def _add_balance_summary(doc: Document, party: PartySettlement) -> None:
    balance_label = (
        "Nachzahlung" if party.balance_type == "nachzahlung"
        else "Guthaben" if party.balance_type == "guthaben"
        else "Ausgleich"
    )
    rows = [
        ("Summe Nebenkosten", party.total_costs, True),
        ("Abzüglich Vorauszahlungen", party.total_advance_payments, False),
        (balance_label, abs(party.balance), True),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    set_table_no_borders(table)
    set_table_column_widths(table, [5.5, 2.5])

    for row_idx, (label, amount, bold) in enumerate(rows):
        label_cell = table.rows[row_idx].cells[0]
        amount_cell = table.rows[row_idx].cells[1]
        size = 9 if bold and row_idx == len(rows) - 1 else 8
        set_cell_margins(label_cell, top=6, bottom=6, left=0, right=40)
        set_cell_margins(amount_cell, top=6, bottom=6, left=0, right=0)
        set_cell_text(label_cell, f"{label}:", bold=bold, size=size, align_right=True, compact=True)
        set_cell_text(amount_cell, format_eur(amount), bold=bold, size=size, align_right=True, compact=True)


def generate_settlement_docx(
    *,
    preview: SettlementPreview,
    party: PartySettlement,
    landlord_name: str,
    landlord_street: str,
    landlord_city: str,
    landlord_phone: str,
    landlord_email: str,
    apartment_street: str,
    apartment_city: str,
    apartment_iban: str,
    apartment_account_holder: str,
    payment_reference_hint: str,
    payment_text_template: str,
    logo_path: str | None = None,
    person_period_lines: list[PersonPeriodLine] | None = None,
) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    header = doc.add_table(rows=1, cols=2)
    header.autofit = False
    set_table_no_borders(header)
    set_table_column_widths(header, [4.0, 12.5])
    left = header.rows[0].cells[0]
    right = header.rows[0].cells[1]
    if logo_path:
        left.paragraphs[0].add_run().add_picture(logo_path, width=Cm(2.5))
    else:
        set_cell_text(left, "")

    set_cell_text(right, landlord_name, bold=True, size=12, align_right=True, compact=True)
    add_cell_paragraph(right, landlord_street, size=9, align_right=True)
    add_cell_paragraph(right, landlord_city, size=9, align_right=True)
    if landlord_phone:
        add_cell_paragraph(right, f"Tel. {landlord_phone}", size=9, align_right=True)
    if landlord_email:
        add_cell_paragraph(right, landlord_email, size=8, align_right=True)

    add_styled_paragraph(doc, party.tenant_name, size=10, compact=True)
    add_styled_paragraph(
        doc,
        f"{party.room_name} · {apartment_street} · {apartment_city}",
        size=9,
        compact=True,
    )

    add_styled_paragraph(
        doc,
        f"Betriebskostenabrechnung {preview.year}",
        bold=True,
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        compact=True,
    )
    add_styled_paragraph(
        doc,
        f"Sehr geehrte/r {party.tenant_name}, anbei die Nebenkostenabrechnung "
        f"für 01.01.{preview.year}–31.12.{preview.year}.",
        size=9,
        compact=True,
    )

    _add_allocation_notes(
        doc,
        preview=preview,
        party=party,
        person_period_lines=person_period_lines or [],
    )

    _add_cost_lines_table(doc, party=party)

    _add_balance_summary(doc, party)

    payment_text = payment_text_template.strip() or _default_payment_text(
        apartment_iban, apartment_account_holder, payment_reference_hint
    )
    add_styled_paragraph(doc, payment_text, size=8, compact=True)

    return doc


def settlement_docx_bytes(doc: Document) -> bytes:
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
